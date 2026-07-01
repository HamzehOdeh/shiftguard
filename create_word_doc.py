"""Generate the full business plan as a formatted Word document."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- TITLE PAGE ---
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Workforce Compliance AI')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Complete Business & Technical Plan')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('AI-powered schedule compliance for hourly, shift-based,\ndistributed workforces')
run.font.size = Pt(14)
run.italic = True

for _ in range(6):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Prepared by: Hamzeh Hodeh\n').bold = True
info.add_run('July 2026\n')
info.add_run('CONFIDENTIAL')

doc.add_page_break()

# --- TABLE OF CONTENTS ---
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. The Market Opportunity',
    '3. Product Overview',
    '4. Competitive Landscape',
    '5. Technical Architecture',
    '6. Go-to-Market Strategy',
    '7. 90-Day Validation Plan',
    '8. Financial Projections',
    '9. Deployment Steps (This Week)',
    '10. IP & Legal Considerations',
    '11. Exit Scenarios',
    '12. Outreach Templates',
    '13. Next Steps',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# --- SECTION 1: EXECUTIVE SUMMARY ---
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph('What: ', style='List Bullet').add_run('An AI-powered compliance layer that checks shift schedules against labor laws, union contracts, and company policies — catching violations before they cost money.')
doc.add_paragraph('For whom: ', style='List Bullet').add_run('Any company with 500+ hourly workers (warehouses, logistics, retail, healthcare, hospitality, manufacturing).')
doc.add_paragraph('The problem: ', style='List Bullet').add_run('80M+ hourly workers in the US. Predictive scheduling laws are exploding (12+ states). No existing tool does AI-native compliance reasoning. Violations discovered after the fact cost $5K-$50K+ per quarter.')
doc.add_paragraph('The solution: ', style='List Bullet').add_run('Upload a schedule → get instant violations + compliant alternatives + cost impact. One second vs. hours of manual checking.')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Your Unfair Advantage: ').bold = True
p.add_run('5 years managing millions of labor hours at Amazon Ops Finance. You\'ve lived the pain of labor law + union rules + operational needs colliding daily. Most YC founders trying to build in this space have never met an hourly worker.')

doc.add_page_break()

# --- SECTION 2: MARKET OPPORTUNITY ---
doc.add_heading('2. The Market Opportunity', level=1)

doc.add_heading('Problem Statement', level=2)
doc.add_paragraph('When a shift scheduler builds next week\'s schedule, they must simultaneously comply with:')
doc.add_paragraph('Federal labor law (FMLA, FLSA)', style='List Bullet')
doc.add_paragraph('State labor law (varies by state - predictive scheduling, meal breaks, overtime)', style='List Bullet')
doc.add_paragraph('City/county ordinances (Chicago, NYC, Oregon, Seattle, etc.)', style='List Bullet')
doc.add_paragraph('Union collective bargaining agreements (CBAs)', style='List Bullet')
doc.add_paragraph('Company internal policies (hour caps, minor restrictions, training)', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Current state: ').bold = True
p.add_run('Managers use tribal knowledge, get lucky, or find out weeks later when a grievance is filed or a fine arrives.')

doc.add_heading('Cost of Failure', level=2)
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Shading Accent 1'
table.cell(0, 0).text = 'Violation Type'
table.cell(0, 1).text = 'Cost'
table.cell(1, 0).text = 'Chicago Fair Workweek'
table.cell(1, 1).text = '$300-$500 per employee per violation'
table.cell(2, 0).text = 'California meal break'
table.cell(2, 1).text = '1 hour premium pay per violation per day'
table.cell(3, 0).text = 'FMLA interference'
table.cell(3, 1).text = '$50K-$500K+ in damages'
table.cell(4, 0).text = 'Union grievances'
table.cell(4, 1).text = '$5K-$50K per grievance'
table.cell(5, 0).text = 'Aggregate annual'
table.cell(5, 1).text = '$50K-$500K+ per year for mid-size operation'

doc.add_heading('Market Size', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Shading Accent 1'
table.cell(0, 0).text = 'Tier'
table.cell(0, 1).text = 'Size'
table.cell(0, 2).text = 'Calculation'
table.cell(1, 0).text = 'TAM'
table.cell(1, 1).text = '$2.7B'
table.cell(1, 2).text = '18,000 US companies x $150K ACV'
table.cell(2, 0).text = 'SAM'
table.cell(2, 1).text = '$1.08B'
table.cell(2, 2).text = '7,200 companies in regulated states'
table.cell(3, 0).text = 'SOM (Year 1)'
table.cell(3, 1).text = '$1.6M'
table.cell(3, 2).text = '15 early-adopter customers x $105K ACV'

doc.add_heading('Regulatory Tailwind', level=2)
doc.add_paragraph('States/cities with predictive scheduling laws (growing 2-3 per year):')
states = ['Oregon (2018)', 'New York City (2017)', 'Chicago (2020)', 'Seattle (2017)',
          'San Francisco (2015)', 'Philadelphia (2020)', 'Los Angeles (2023)',
          'New Jersey (proposed)', 'Massachusetts (proposed)', 'Connecticut (proposed)']
for s in states:
    doc.add_paragraph(s, style='List Bullet')

doc.add_page_break()

# --- SECTION 3: PRODUCT OVERVIEW ---
doc.add_heading('3. Product Overview', level=1)

doc.add_heading('How It Works', level=2)
doc.add_paragraph('INPUT: Shift schedule (CSV/Excel/API from Kronos/UKG) + Employee data + Jurisdiction')
doc.add_paragraph('PROCESSING: AI checks every shift against all applicable rules in <1 second')
doc.add_paragraph('OUTPUT: Violations with severity, legal citation, cost impact, and recommended fix')

doc.add_heading('Demo Results (Current Prototype)', level=2)
doc.add_paragraph('The prototype analyzes a 10-employee, 52-shift warehouse schedule and catches 15 violations:')

table = doc.add_table(rows=10, cols=4)
table.style = 'Light Shading Accent 1'
headers = ['#', 'Severity', 'Violation', 'Cost Impact']
for i, h in enumerate(headers):
    table.cell(0, i).text = h
data = [
    ('1', 'CRITICAL', 'Minor scheduled past 10pm', 'Legal fine + termination risk'),
    ('2', 'CRITICAL', '5 shifts during FMLA leave', '$50K-$500K claim'),
    ('3', 'HIGH', '14-day notice not met (Chicago)', '$3,000-$5,000 penalty'),
    ('4', 'HIGH', 'Clopening: 7.5hrs between shifts', '1.25x pay + grievance'),
    ('5', 'HIGH', '8 consecutive days (max 6)', 'Double time + grievance'),
    ('6', 'HIGH', '3-hour shift (min 4 per CBA)', 'Pay for full 4 hours'),
    ('7', 'MEDIUM', 'CBA posting deadline missed', 'Grievance risk'),
    ('8', 'MEDIUM', '63 weekly hours (max 60)', 'VP approval required'),
    ('9', 'MEDIUM', 'PTO scheduling conflicts', 'Morale + grievance'),
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        table.cell(row_idx, col_idx).text = val

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Total exposure: $4,500 - $8,000+ for ONE week at ONE facility.').bold = True

doc.add_heading('Rule Categories Covered', level=2)
categories = [
    'Predictive Scheduling Laws (Oregon, NYC, Chicago, Seattle, LA)',
    'Meal & Rest Break Laws (California, Oregon, Washington)',
    'Union CBA Rules (configurable per contract)',
    'Leave & Time-Off (Federal FMLA + State sick leave + PTO)',
    'Company Policy (configurable per organization)',
]
for c in categories:
    doc.add_paragraph(c, style='List Bullet')

doc.add_page_break()

# --- SECTION 4: COMPETITIVE LANDSCAPE ---
doc.add_heading('4. Competitive Landscape', level=1)

table = doc.add_table(rows=6, cols=4)
table.style = 'Light Shading Accent 1'
headers = ['Player', 'Revenue', 'What they do', 'Gap']
for i, h in enumerate(headers):
    table.cell(0, i).text = h
competitors = [
    ('UKG (Kronos)', '$6B+', 'Scheduling, time tracking', 'Rules take 3-6 months to update'),
    ('Workday', '$7B+', 'HCM for salaried', 'Not built for hourly complexity'),
    ('Legion', '$100M+', 'AI scheduling optimization', 'Shallow compliance'),
    ('Deputy', '$100M+', 'SMB scheduling', 'No enterprise compliance'),
    ('ADP', '$18B+', 'Payroll, HR', 'Post-hoc, not preventive'),
]
for row_idx, row_data in enumerate(competitors, 1):
    for col_idx, val in enumerate(row_data):
        table.cell(row_idx, col_idx).text = val

doc.add_heading('Your Differentiator', level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Shading Accent 1'
table.cell(0, 0).text = 'Dimension'
table.cell(0, 1).text = 'Incumbents'
table.cell(0, 2).text = 'You'
diffs = [
    ('Rule update speed', '3-6 months', '48 hours (AI parses new laws)'),
    ('CBA encoding', 'Professional services ($50K+)', 'Self-service AI ingestion'),
    ('Approach', 'Reactive (after violation)', 'Preventive (before publishing)'),
    ('Deployment', 'Rip-and-replace (12-18mo)', 'Layer on top of existing tools'),
    ('Cost', '$500K+ implementation', '$15-25 PEPM, no implementation'),
]
for row_idx, row_data in enumerate(diffs, 1):
    for col_idx, val in enumerate(row_data):
        table.cell(row_idx, col_idx).text = val

doc.add_page_break()

# --- SECTION 5: TECHNICAL ARCHITECTURE ---
doc.add_heading('5. Technical Architecture', level=1)

doc.add_heading('Production Stack', level=2)
table = doc.add_table(rows=8, cols=3)
table.style = 'Light Shading Accent 1'
table.cell(0, 0).text = 'Component'
table.cell(0, 1).text = 'Tool'
table.cell(0, 2).text = 'Cost'
stack = [
    ('Frontend', 'React or Streamlit', 'Free'),
    ('Backend API', 'FastAPI (Python)', 'Free'),
    ('Database', 'PostgreSQL (Supabase)', '$0-25/mo'),
    ('Hosting', 'AWS / Railway', '$20-100/mo'),
    ('Auth', 'Clerk or Auth0', 'Free tier'),
    ('AI Reasoning', 'Claude API (Anthropic)', '$20-200/mo'),
    ('Compliance Data', 'LegiScan + scrapers + AI', '$200-500/mo'),
]
for row_idx, row_data in enumerate(stack, 1):
    for col_idx, val in enumerate(row_data):
        table.cell(row_idx, col_idx).text = val

doc.add_heading('The AI Parsing Moat', level=2)
doc.add_paragraph('When a new law passes:')
doc.add_paragraph('1. LegiScan alerts you (or scrape state legislature site)')
doc.add_paragraph('2. Feed law text to Claude: "Parse into structured compliance rules"')
doc.add_paragraph('3. Claude outputs structured rules (ID, requirement, penalty, severity)')
doc.add_paragraph('4. You review for accuracy (30 minutes)')
doc.add_paragraph('5. Deploy to all affected customers')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Result: New law → live in product in 48 hours. Incumbents: 3-6 months.').bold = True

doc.add_page_break()

# --- SECTION 6: GO-TO-MARKET ---
doc.add_heading('6. Go-to-Market Strategy', level=1)

doc.add_heading('Pricing Model', level=2)
table = doc.add_table(rows=4, cols=4)
table.style = 'Light Shading Accent 1'
table.cell(0, 0).text = 'Tier'
table.cell(0, 1).text = 'Price'
table.cell(0, 2).text = 'Target'
table.cell(0, 3).text = 'Includes'
pricing = [
    ('Starter', '$15/emp/month', '500-2,000 employees', 'State + federal + basic CBA'),
    ('Professional', '$22/emp/month', '2,000-10,000 employees', 'Multi-state, full CBA, AI'),
    ('Enterprise', '$25/emp/month', '10,000+ employees', 'Custom rules, API, support'),
]
for row_idx, row_data in enumerate(pricing, 1):
    for col_idx, val in enumerate(row_data):
        table.cell(row_idx, col_idx).text = val

doc.add_heading('Ideal First Customers', level=2)
customers = [
    '3PL warehouses (XPO, Ryder, DHL Supply Chain) — multi-state, union, high risk',
    'Regional grocery chains — Chicago/CA predictive scheduling, unionized',
    'Hospital systems — complex scheduling, high penalty stakes',
    'Food distribution (Sysco, US Foods) — warehouse + delivery, multi-state',
    'Manufacturing (automotive, food processing) — union + state compliance',
]
for c in customers:
    doc.add_paragraph(c, style='List Bullet')

doc.add_page_break()

# --- SECTION 7: 90-DAY PLAN ---
doc.add_heading('7. 90-Day Validation Plan', level=1)

doc.add_heading('Weeks 1-4: Discovery Sprint', level=2)
doc.add_paragraph('Goal: Complete 20 discovery interviews. Validate problem severity.')
items = [
    'Build target list of 50 companies (warehouse, logistics, retail)',
    'Send 100 outreach messages (target 20% response rate)',
    'Complete 20 discovery interviews',
    'Document: pain score (1-10), willingness to pay, current tools, budget timing',
    'Identify 3-5 design partners willing to co-develop',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Weeks 5-8: Solution Validation', level=2)
doc.add_paragraph('Goal: Build MVP scope from data. Get 3 Letters of Intent.')
items = [
    'Synthesize interviews into Jobs to Be Done framework',
    'Demo prototype to top 10 interviewees (target NPS > 8/10)',
    'Validate pricing: $15-25 PEPM (target > 50% acceptance)',
    'Secure 3-5 Letters of Intent',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Weeks 9-12: Go/No-Go Decision', level=2)
p = doc.add_paragraph()
p.add_run('GO if: ').bold = True
p.add_run('≥ 3 LOIs, pain score ≥ 8/10, ACV ≥ $100K, clear differentiator confirmed')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('PIVOT if: ').bold = True
p.add_run('Pain real but pricing won\'t support venture scale → consulting approach')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('STOP if: ').bold = True
p.add_run('Pain score < 6/10, < 2 LOIs after 50 conversations')

doc.add_page_break()

# --- SECTION 8: FINANCIALS ---
doc.add_heading('8. Financial Projections', level=1)

doc.add_heading('Year 1 Revenue', level=2)
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Shading Accent 1'
table.cell(0, 0).text = 'Quarter'
table.cell(0, 1).text = 'Customers'
table.cell(0, 2).text = 'Revenue'
table.cell(0, 3).text = 'Notes'
rev = [
    ('Q1', '2 (pilots)', '$30K', 'Consulting-style'),
    ('Q2', '5', '$90K', 'Convert + 3 new'),
    ('Q3', '10', '$180K', 'Self-serve onboarding'),
    ('Q4', '15', '$270K', 'First AE hired'),
]
for row_idx, row_data in enumerate(rev, 1):
    for col_idx, val in enumerate(row_data):
        table.cell(row_idx, col_idx).text = val

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Year 1 Total: $570K ARR | Year 2 Target: $2.5M ARR').bold = True

doc.add_heading('Cost Structure (Year 1)', level=2)
table = doc.add_table(rows=8, cols=3)
table.style = 'Light Shading Accent 1'
table.cell(0, 0).text = 'Item'
table.cell(0, 1).text = 'Monthly'
table.cell(0, 2).text = 'Annual'
costs = [
    ('Salary (deferred/low)', '$5K', '$60K'),
    ('Cloud hosting', '$200', '$2.4K'),
    ('Claude API', '$200', '$2.4K'),
    ('Data sources (LegiScan etc)', '$300', '$3.6K'),
    ('Legal (entity + contracts)', '-', '$5K'),
    ('Sales tools', '$200', '$2.4K'),
    ('TOTAL', '$5.9K', '$75.8K'),
]
for row_idx, row_data in enumerate(costs, 1):
    for col_idx, val in enumerate(row_data):
        table.cell(row_idx, col_idx).text = val

doc.add_paragraph()
doc.add_paragraph('Breakeven: ~5 customers at $15K ACV = $75K')

doc.add_page_break()

# --- SECTION 9: DEPLOYMENT ---
doc.add_heading('9. Deployment Steps (This Week)', level=1)

steps = [
    ('Step 1: GitHub Setup (15 min)', 'Create account at github.com/signup. Create private repo "workforce-compliance-ai". Push code.'),
    ('Step 2: Streamlit Cloud (10 min)', 'Go to share.streamlit.io. Sign in with GitHub. Deploy app.py. Get live URL.'),
    ('Step 3: Claude API Key (5 min)', 'Go to console.anthropic.com. Create account. Generate API key. Add to Streamlit secrets.'),
    ('Step 4: Test & Share', 'Open deployed URL. Run demo. Share link in customer discovery calls.'),
]
for title, desc in steps:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# --- SECTION 10: IP & LEGAL ---
doc.add_heading('10. IP & Legal Considerations', level=1)

doc.add_heading('While Still at Amazon', level=2)
doc.add_paragraph('DO: Build on your own time, own equipment, own accounts', style='List Bullet')
doc.add_paragraph('DO: Use publicly available legal information (laws are public)', style='List Bullet')
doc.add_paragraph('DO: Talk to people outside Amazon for customer discovery', style='List Bullet')
doc.add_paragraph("DON'T: Use any Amazon proprietary data, tools, or code", style='List Bullet')
doc.add_paragraph("DON'T: Use Amazon email or systems for this project", style='List Bullet')
doc.add_paragraph("DON'T: Solicit Amazon as a customer while employed", style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key principle: ').bold = True
p.add_run('Your knowledge of how scheduling works is yours. Amazon\'s data and tools are theirs.')

doc.add_page_break()

# --- SECTION 11: EXIT ---
doc.add_heading('11. Exit Scenarios', level=1)

table = doc.add_table(rows=7, cols=3)
table.style = 'Light Shading Accent 1'
table.cell(0, 0).text = 'Acquirer'
table.cell(0, 1).text = 'Valuation Range'
table.cell(0, 2).text = 'Why'
exits = [
    ('UKG', '$200M-$1B+', 'AI compliance for 50K+ customers'),
    ('Workday', '$100M-$500M', 'Enter hourly workforce market'),
    ('ADP', '$200M-$1B+', 'Compliance + payroll integration'),
    ('Ceridian/Dayforce', '$100M-$500M', 'Differentiate vs UKG'),
    ('ServiceNow', '$200M-$1B+', 'Extend HR workflow automation'),
    ('Private Equity', '$50M-$200M', 'Roll up workforce mgmt'),
]
for row_idx, row_data in enumerate(exits, 1):
    for col_idx, val in enumerate(row_data):
        table.cell(row_idx, col_idx).text = val

doc.add_page_break()

# --- SECTION 12: OUTREACH ---
doc.add_heading('12. Outreach Templates', level=1)

doc.add_heading('Cold Outreach to VP Ops', level=2)
doc.add_paragraph('Subject: Quick question about scheduling compliance at {company_name}')
doc.add_paragraph()
doc.add_paragraph('Hi {first_name},')
doc.add_paragraph()
doc.add_paragraph('I noticed {company_name} operates distribution centers in {state} — which means you\'re navigating the {law_name} predictive scheduling requirements.')
doc.add_paragraph()
doc.add_paragraph('I spent 5 years in Ops Finance at Amazon building workforce planning models, and I\'m now researching how operations leaders handle the compliance complexity of multi-state scheduling laws. I\'m not selling anything — I\'m interviewing 15-20 ops leaders to validate whether this problem is worth solving with better tooling.')
doc.add_paragraph()
doc.add_paragraph('Would you have 25 minutes this week or next? I\'d love to hear how your team handles it today.')
doc.add_paragraph()
doc.add_paragraph('Best,\nHamzeh')

doc.add_heading('Warm Intro via Amazon Alumni', level=2)
doc.add_paragraph('Subject: Fellow Amazonian exploring workforce compliance — intro request')
doc.add_paragraph()
doc.add_paragraph('Hi {mutual_connection},')
doc.add_paragraph()
doc.add_paragraph('I\'m building in the workforce compliance space (AI that ensures shift schedules comply with predictive scheduling laws and union rules before they\'re published). Think "compliance guardrails for scheduling."')
doc.add_paragraph()
doc.add_paragraph('I noticed you\'re connected to {target_name} at {company_name}. Would you be open to making a quick intro? I\'m in discovery mode — just learning how their team handles multi-state scheduling compliance today.')
doc.add_paragraph()
doc.add_paragraph('Happy to return the favor. Thanks!\nHamzeh')

doc.add_page_break()

# --- SECTION 13: NEXT STEPS ---
doc.add_heading('13. Next Steps (Your Action Items)', level=1)

actions = [
    ('TODAY', 'Create GitHub account + Anthropic API account'),
    ('THIS WEEK', 'Deploy app to Streamlit Cloud. Send first 10 outreach messages.'),
    ('NEXT 2 WEEKS', 'Complete 5 discovery interviews using interview framework.'),
    ('30 DAYS', 'Prototype demoed to 10 people. 2-3 design partner candidates.'),
    ('60 DAYS', 'First LOI signed or clear pivot signal.'),
    ('90 DAYS', 'Go/No-Go decision with data to back it up.'),
]

table = doc.add_table(rows=7, cols=2)
table.style = 'Light Shading Accent 1'
table.cell(0, 0).text = 'Timeline'
table.cell(0, 1).text = 'Action'
for row_idx, (timeline, action) in enumerate(actions, 1):
    table.cell(row_idx, 0).text = timeline
    table.cell(row_idx, 1).text = action

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('— END OF DOCUMENT —').italic = True

# Save
output_path = r'C:\Users\hodeh\Documents\Support Call\workforce-compliance-ai\Workforce_Compliance_AI_Plan.docx'
doc.save(output_path)
print(f'Word document saved: {output_path}')
